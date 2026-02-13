from __future__ import annotations

import re
from io import BytesIO
from typing import List

from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn

from epub_processing import Article


def _set_font_chinese_english(font, chinese_font: str, english_font: str):
    """设置字体的中文字体和英文字体。"""
    font.name = english_font
    # 通过XML直接设置中文字体
    r_fonts = font._element.find(qn('w:rFonts'))
    if r_fonts is None:
        from docx.oxml import parse_xml
        r_fonts = parse_xml(r'<w:rFonts xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>')
        font._element.append(r_fonts)
    r_fonts.set(qn('w:eastAsia'), chinese_font)  # 中文字体
    r_fonts.set(qn('w:ascii'), english_font)  # 英文字体
    r_fonts.set(qn('w:hAnsi'), english_font)  # 高ANSI字符字体


def _extract_article_title(title: str) -> str:
    """从标题中提取真正的文章标题，去除网址和前缀，支持多行。"""
    # 先统一处理换行符，合并为单行以便正则匹配
    title_normalized = re.sub(r'\s+', ' ', title).strip()
    
    # 先去除网址模式（feed_*/article_*/index_*.html 等）
    title_normalized = re.sub(r'feed_\d+/article_\d+/index_\w+\.html', '', title_normalized)
    title_normalized = re.sub(r'[a-zA-Z0-9_/]+\.html', '', title_normalized)
    
    # 优先提取《》中的内容（支持跨行）
    match = re.search(r'《([^》]+)》', title_normalized)
    if match:
        return match.group(1).strip()
    
    # 提取【文章标题】** :或【文章标题】：后面的内容（支持跨行）
    match = re.search(r'【文章标题】\*?\*?\s*[：:]\s*(.+)', title_normalized)
    if match:
        extracted = match.group(1).strip()
        # 如果提取的内容还包含《》，提取《》中的内容
        inner_match = re.search(r'《([^》]+)》', extracted)
        if inner_match:
            return inner_match.group(1).strip()
        return extracted
    
    # 去除【文章标题】等前缀标记
    title_normalized = re.sub(r'【文章标题】\*+\s*[：:]\s*', '', title_normalized)
    title_normalized = title_normalized.strip()
    
    # 如果还有残留的网址或路径，继续清理
    title_normalized = re.sub(r'^[a-zA-Z0-9_/\.]+', '', title_normalized).strip()
    
    # 如果清理后还有内容，返回清理后的内容
    if title_normalized:
        return title_normalized
    
    return "未命名文章"


# 刊名等通用标题，若提取到这些则视为无效，需从引言兜底
_GENERIC_TITLES = frozenset({"经济学人", "the economist", "economist", "《经济学人》"})

# 无意义的占位标题（如「标题」），不得作为口播稿/分析标题使用
_INVALID_TITLE_VALUES = frozenset({"标题", "title"})


def _derive_title_from_intro(analysis: str) -> str | None:
    """从【引言】段落提取首句或前若干字作为兜底标题。综述类引言优先提取简短栏目名。"""
    match = re.search(r'【引言】\*?\*?\s*[：:]\s*(.+?)(?=\n\n|【|$)', analysis, re.DOTALL)
    if not match:
        return None
    intro = match.group(1).strip()
    if not intro:
        return None
    intro_start = intro[:80]
    # 综述类引言：提取简短栏目名
    if re.search(r"全球商业", intro_start):
        return "全球商业动态"
    if re.search(r"全球市场", intro_start):
        return "全球市场动态"
    if re.search(r"全球政治", intro_start):
        return "全球政治动态"
    if re.search(r"本周全球", intro_start):
        return "本周全球商业动态"
    # 取首句（以。！？为界）或前 40 字
    for sep in "。", "！", "？", "，":
        idx = intro.find(sep)
        if idx > 0 and idx <= 50:
            return intro[: idx + 1].strip()
    return intro[:40].strip() + ("…" if len(intro) > 40 else "") if intro else None


def _is_roundup_intro(analysis: str) -> bool:
    """判断引言是否为综述类（本周全球、全球市场等）。"""
    match = re.search(r'【引言】\*?\*?\s*[：:]\s*(.+?)(?=\n\n|【|$)', analysis, re.DOTALL)
    if not match:
        return False
    intro = (match.group(1) or "")[:80]
    return bool(re.search(r"本周全球|全球市场|全球商业|全球政治", intro))


def _is_specific_topic_title(title: str) -> bool:
    """判断标题是否为具体议题（如法案、政策名），易与综述类文章混淆。"""
    if not title or len(title) < 2:
        return False
    return bool(re.search(r"法案$|法$", title))


def _extract_title_from_analysis(analysis: str) -> str | None:
    """从 DeepSeek 分析结果中提取【文章标题】后的中文标题。"""
    if not analysis or not analysis.strip():
        return None
    # 匹配【文章标题】**：或【文章标题】：后面的内容（可能跨行）
    match = re.search(r'【文章标题】\*?\*?\s*[：:]\s*(.+)', analysis, re.DOTALL)
    if not match:
        return None
    extracted = match.group(1).strip()
    # 优先取《》中的内容
    inner = re.search(r'《([^》]+)》', extracted)
    if inner:
        first_line = inner.group(1).strip()
    else:
        first_line = extracted.split('\n')[0].strip()
    if not first_line:
        return None
    # 若为刊名等通用标题，尝试从引言兜底
    if first_line in _GENERIC_TITLES or first_line.lower() in _GENERIC_TITLES:
        return _derive_title_from_intro(analysis)
    # 综述类引言 + 具体议题式标题（如「叛乱法」「就业权利法案」）：强制使用引言兜底
    if _is_specific_topic_title(first_line) and _is_roundup_intro(analysis):
        derived = _derive_title_from_intro(analysis)
        if derived:
            return derived
        # 如果 _derive_title_from_intro 返回 None，仍然不应该使用具体议题式标题
        # 尝试从引言中提取首句作为兜底
        intro_match = re.search(r'【引言】\*?\*?\s*[：:]\s*(.+?)(?=\n\n|【|$)', analysis, re.DOTALL)
        if intro_match:
            intro = intro_match.group(1).strip()
            # 提取首句（以。！？为界）或前 40 字
            for sep in "。", "！", "？":
                idx = intro.find(sep)
                if idx > 0 and idx <= 50:
                    return intro[: idx + 1].strip()
            if intro:
                return intro[:40].strip() + ("…" if len(intro) > 40 else "")
    return first_line


# 口播稿常见开头/过渡语，从正文开头去掉（标题推导与正文均用）
_AUDIO_SCRIPT_OPENERS = re.compile(
    r"^(?:\s*"
    r"(?:好的,?请听这篇来自《经济学人》的文章。"
    r"|好的,?请听《经济学人》中文(?:版)?有声书。?(?:今天是第\d+篇,?共\d+篇。?)?"
    r"|这是《经济学人》中文(?:版)?有声书,?今天[^。]*。"
    r"|这是《经济学人》中文(?:版)?有声书[，,]\s*第\d+篇[，,]\s*共\d+篇。?"
    r"|这是第\d+篇[，,]\s*共\d+篇。?"
    r"|我们为您播报[^。]*。"
    r"|欢迎收听《经济学人》读者来信栏目。[^。]*。(?:今天是[^。]+。)?(?:本周我们选取的来信[^。]*。)?"
    r")\s*)*",
    re.IGNORECASE,
)


def _derive_title_from_analysis_body(analysis: str) -> str | None:
    """从口播稿正文去掉常见开头后取首句作为标题兜底。若首行为无意义「标题」则跳过，用下一行。"""
    if not analysis or not analysis.strip():
        return None
    rest = _AUDIO_SCRIPT_OPENERS.sub("", analysis.strip()).strip()
    if not rest:
        return None
    lines = [ln.strip() for ln in rest.split("\n") if ln.strip()]
    if not lines:
        return None
    # 若首行仅为「标题」或「**标题**」，丢弃该行，用下一行作为候选
    first_clean = re.sub(r"^\*+|\*+$", "", lines[0]).strip()
    if first_clean == "标题" and len(lines) > 1:
        rest = "\n".join(lines[1:])
    else:
        rest = "\n".join(lines)
    if not rest.strip():
        return None
    for sep in "。", "？", "！":
        idx = rest.find(sep)
        if idx >= 0:
            first_sentence = rest[: idx + 1].strip()
            if len(first_sentence) > 40:
                first_sentence = first_sentence[:40] + "…"
            return first_sentence if first_sentence else None
    title = rest[:40].strip() + ("…" if len(rest) > 40 else "")
    return title if title else None


def _strip_listen_openers(text: str) -> str:
    """去掉口播稿开头的过渡语，返回从正文开始的内容。"""
    if not text or not text.strip():
        return text
    return _AUDIO_SCRIPT_OPENERS.sub("", text.strip()).strip()


# 口播稿结尾需删除的推广/下载类句子（从文末循环剥离）
_LISTEN_CLOSING_PATTERNS = [
    re.compile(r"\s*要了解[^。]+《欧洲咖啡馆》[^。]*。\s*$"),
    re.compile(r"\s*这篇文章由\s*calibre\s*从[^。]*。?\s*$", re.IGNORECASE),
    re.compile(r"\s*请订阅[^。]*。\s*$"),
    re.compile(r"\s*从以下[^。]*网址[^。]*下载[^。]*。?\s*$"),
]


def _strip_listen_closings(text: str) -> str:
    """去掉口播稿结尾的推广/下载类句子，返回清理后的内容。"""
    if not text or not text.strip():
        return text
    body = text.strip()
    changed = True
    while changed:
        changed = False
        for pat in _LISTEN_CLOSING_PATTERNS:
            new_body = pat.sub("", body).strip()
            if new_body != body:
                body = new_body
                changed = True
    return body


def build_docx_from_analyses(
    analyses: List[str],
    articles: List[Article],
    titles_override: List[str] | None = None,
) -> BytesIO:
    """将所有文章的中文分析结果写入单一 Word 文档并返回内存流。
    titles_override: 若提供且与 analyses 等长，则优先用其非空项作为标题（与「看我」一致）。"""
    if len(analyses) != len(articles):
        raise ValueError("analyses 与 articles 数量不一致。")
    if titles_override is not None and len(titles_override) != len(articles):
        raise ValueError("titles_override 与 articles 数量不一致。")

    doc = Document()

    # 设置全局字体：中文使用微软雅黑，英文使用Times New Roman
    style = doc.styles["Normal"]
    font = style.font
    _set_font_chinese_english(font, "微软雅黑", "Times New Roman")
    font.size = Pt(11)
    
    # 设置标题字体：中文使用微软雅黑，英文使用Times New Roman
    heading_style = doc.styles["Heading 1"]
    heading_font = heading_style.font
    _set_font_chinese_english(heading_font, "微软雅黑", "Times New Roman")

    for i, (article, analysis) in enumerate(zip(articles, analyses, strict=True)):
        # 优先用「看我」翻译标题，否则从分析/EPUB/口播稿正文推导
        from_override = (titles_override[i] or "").strip() if titles_override else ""
        if from_override and from_override != "未命名文章":
            heading = from_override
        else:
            from_epub = _extract_article_title(article.title)
            heading = (
                _extract_title_from_analysis(analysis)
                or (from_epub if (from_epub and from_epub != "未命名文章" and from_epub.strip().lower() not in _INVALID_TITLE_VALUES) else None)
                or _derive_title_from_analysis_body(analysis)
                or "未命名文章"
            )
        # 规范化：去掉首尾 *、去掉「标题：」前缀，得到纯标题文字后统一为「标题：XX」格式
        heading = re.sub(r"^\*+|\*+$", "", heading).strip()
        heading = re.sub(r"^#?\s*标题\s*[：:]\s*", "", heading).strip()
        if not heading or heading.strip().lower() in _INVALID_TITLE_VALUES:
            heading = "未命名文章"
        if not heading.startswith("标题：") and not heading.startswith("标题:"):
            heading = "标题：" + heading
        doc.add_heading(heading, level=1)
        # 用于正文去重：仅标题文字（不含「标题：」前缀）
        title_content = re.sub(r"^\s*标题\s*[：:]\s*", "", heading).strip() or heading

        # 去掉口播稿末尾固定结束语与推广/下载句，再去掉开头过渡语，然后拆段
        body_text = re.sub(r"\s*这篇文章就为您播报到这里。感谢您的收听。\s*$", "", analysis)
        body_text = _strip_listen_closings(body_text)
        body_text = _strip_listen_openers(body_text)
        chunks = [chunk.strip() for chunk in body_text.split("\n\n") if chunk.strip()]
        if not chunks:
            continue

        # 匹配「仅标题标签」的整行（用于丢弃）；匹配段落开头的标题前缀（用于剥离）
        _title_only_line = re.compile(r"^\s*(?:#+\s*|\|\s*|\*+\s*)*标题\s*[：:]*\s*$", re.IGNORECASE)
        _title_label_prefix = re.compile(r"^\s*(?:#+\s*|\|\s*|\*+\s*)*标题\s*[：:]*\s*", re.IGNORECASE)
        for chunk in chunks:
            # 跳过【文章标题】段落
            if re.match(r'^【文章标题】\*?\*?\s*[：:]', chunk):
                continue
            # 跳过仅含「标题」或「**标题**」的无意义字段行（整段只有标签时跳过）
            if re.match(r'^\s*\*?\*?\s*标题\s*\*?\*?\s*$', chunk):
                continue
            # 先去掉开头多行「仅标题标签」行（如 标题、#### 标题:、*###**标题: 等单独成行）
            lines = chunk.split("\n")
            while lines and _title_only_line.match(lines[0].strip()):
                lines.pop(0)
            chunk = "\n".join(lines).strip()
            # 再剥离首行开头的「标题」/「标题：」及 Markdown 前缀
            chunk = _title_label_prefix.sub("", chunk).strip()
            if not chunk:
                continue
            # 跳过与当前标题重复的段落（首段常为「标题：XX」，已作为 heading 展示）
            if chunk == title_content or chunk.strip() == heading:
                continue
            doc.add_paragraph(chunk)

        # 章节之间空一页或空几行，这里简单空两行
        doc.add_paragraph()
        doc.add_paragraph()

    stream = BytesIO()
    doc.save(stream)
    stream.seek(0)
    return stream


def _parse_translation(translation: str) -> tuple[str, str, str | None]:
    """
    从单条翻译文本解析出：标题、正文、译者注（可选）。
    返回 (title, body, translator_note_or_none)。
    """
    text = (translation or "").strip()
    if not text:
        return ("未命名文章", "", None)

    lines = [ln.strip() for ln in text.splitlines() if ln.strip()]
    if not lines:
        return ("未命名文章", "", None)

    # 标题：首行匹配「标题：」或「【…】」则取该行/括号内容；否则取第一行或前 80 字
    first = lines[0]
    title = "未命名文章"
    title_match = re.match(r"#?\s*标题\s*[：:]\s*(.+)", first)
    if title_match:
        title = title_match.group(1).strip()
        if len(title) > 80:
            title = title[:80]
    else:
        bracket = re.search(r"【([^】]+)】", first)
        if bracket:
            title = bracket.group(1).strip()
        else:
            title = first[:80] if len(first) > 80 else first

    # 正文与译者注：从标题后到「译者注：」之前为正文；若有「译者注：」则单独一段
    rest = "\n\n".join(lines[1:]) if len(lines) > 1 else ""
    translator_note = None
    body = rest

    note_marker = re.search(r"译者注\s*[：:]\s*", rest)
    if note_marker:
        body = rest[: note_marker.start()].strip()
        translator_note = rest[note_marker.end() :].strip()

    title = re.sub(r"^\*+|\*+$", "", title).strip()
    title = re.sub(r"^#?\s*标题\s*[：:]\s*", "", title).strip()
    if not title:
        title = "未命名文章"

    return (title, body, translator_note if translator_note else None)


def build_docx_from_translations(
    translations: List[str],
    articles: List[Article],
) -> BytesIO:
    """将全文翻译结果写入单一 Word 文档并返回内存流。"""
    if len(translations) != len(articles):
        raise ValueError("translations 与 articles 数量不一致。")

    doc = Document()

    style = doc.styles["Normal"]
    font = style.font
    _set_font_chinese_english(font, "微软雅黑", "Times New Roman")
    font.size = Pt(11)

    heading_style = doc.styles["Heading 1"]
    heading_font = heading_style.font
    _set_font_chinese_english(heading_font, "微软雅黑", "Times New Roman")

    for i, (article, raw_translation) in enumerate(zip(articles, translations, strict=True)):
        title, body, translator_note = _parse_translation(raw_translation)
        if title == "未命名文章" and body == "" and not translator_note:
            title = _extract_article_title(article.title) or "未命名文章"

        doc.add_heading(title, level=1)

        for para in body.split("\n\n"):
            para = para.strip()
            if para:
                doc.add_paragraph(para)

        if translator_note:
            doc.add_paragraph()
            doc.add_paragraph("译者注：" + translator_note)

        doc.add_paragraph()
        doc.add_paragraph()

    stream = BytesIO()
    doc.save(stream)
    stream.seek(0)
    return stream

